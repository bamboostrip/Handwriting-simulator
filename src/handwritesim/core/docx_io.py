"""docx 文档解析：提取段落对齐、首行缩进与多角色 Run 流。

增强：
- 读取 <w:rPr><w:highlight w:val="..."/> / <w:shd w:fill="..."/> 背景色，动态映射到角色：
  文档中首次出现的背景色 → 角色2（手写角色1），次出现 → 角色3 … 不限制黄/绿
- 读取 <w:color w:val="..."/> 作为 Run 级颜色覆盖
- 读取 <w:rFonts w:eastAsia/w:ascii> 与 <w:sz> 字体字号，打印体（无背景）沿用原文系统字体，
  手写体则固定用用户选择的手写字体（背景标记段忽略原文纸质字体）；
  Word 字号按文档内比例映射到全局字号坐标系（正文≈全局字号，标题按比例放大）
- 支持文本标签拆分 {{角色名:文本}} / {{手写:文本}} / {{打印:文本}}
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .models import HandwritingRole, Paragraph, TextRun

try:
    from .system_fonts import family_to_file
except Exception:  # 导入失败时回退为空实现
    def family_to_file(family: str):  # type: ignore
        return None

# docx 使用 EMU（1 英寸 = 914400 EMU）；1pt = 12700 EMU（1/72 英寸）
_EMU_PER_PT = 12700

# 标签正则：{{前缀:内容}} 前缀可为 手写/打印/角色1/张三 等任意非冒号非括号字符，内容为任意非 } 字符（非贪婪）
_TAG_RE = re.compile(r"\{\{\s*([^:{}]+?)\s*:\s*(.*?)\s*\}\}", re.DOTALL)


def _first_line_chars(para) -> float | None:
    """读取 w:firstLineChars（1/100 字符），直接格式优先，再沿样式链继承。

    中文 Word 文档的“首行缩进 2 字符”通常写 firstLineChars 而非 EMU，
    python-docx 的 paragraph_format.first_line_indent 读不到它。
    """
    pPr = para._p.pPr
    if pPr is not None:
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            value = ind.get(qn("w:firstLineChars"))
            if value:
                return int(value) / 100.0
    style = para.style
    while style is not None:
        spPr = style.element.find(qn("w:pPr"))
        if spPr is not None:
            ind = spPr.find(qn("w:ind"))
            if ind is not None:
                value = ind.get(qn("w:firstLineChars"))
                if value:
                    return int(value) / 100.0
        style = style.base_style
    return None


def _font_size_pt(para, doc) -> float:
    """取段落实际字号（pt）：run 直接格式 > 段落样式链 > Normal > docDefaults，兜底 12。"""
    for run in para.runs:
        size = run.font.size
        if size:
            return size.pt
    style = para.style
    while style is not None:
        size = style.font.size
        if size:
            return size.pt
        style = style.base_style
    try:
        size = doc.styles["Normal"].font.size
        if size:
            return size.pt
    except KeyError:
        pass
    el = doc.styles.element.find(qn("w:docDefaults"))
    if el is not None:
        el = el.find(qn("w:rPrDefault"))
        if el is not None:
            el = el.find(qn("w:rPr"))
            if el is not None:
                sz = el.find(qn("w:sz"))
                if sz is not None:
                    val = sz.get(qn("w:val"))
                    if val:
                        return int(val) / 2.0  # 半磅 -> pt
    return 12.0


def _first_line_emu_chars(para, doc) -> float | None:
    """把 firstLine（twips→EMU）按文档字号还原为字符数。

    Word/WPS 某些版本把“首行缩进 2 字符”存成固定值 firstLine
    （如 480twips = 12pt 字号的 2 字符）而不写 firstLineChars；
    先按文档字号换算回字符数，再随渲染字号等比缩放，
    否则缩进固定为 96dpi 像素，渲染字号大时明显偏小。
    """
    value = para.paragraph_format.first_line_indent
    if not value:
        style = para.style
        while style is not None:
            value = style.paragraph_format.first_line_indent
            if value:
                break
            style = style.base_style
    if not value:
        return None
    pt = value / _EMU_PER_PT  # EMU -> pt
    return pt / _font_size_pt(para, doc)


def _resolve_align(para) -> str:
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    if para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return "right"
    return "left"


def _run_highlight(run) -> str | None:
    """读取 run 的高亮背景色 w:highlight @w:val（yellow/green/cyan...），或 w:shd @w:fill。

    优先 highlight，其次 shd（Word 底纹）。返回小写色名或 hex（shd），无则 None。
    """
    rPr = run._element.rPr
    if rPr is None:
        return None
    hl = rPr.find(qn("w:highlight"))
    if hl is not None:
        val = hl.get(qn("w:val"))
        if val and val != "none":
            return val.strip().lower()
    shd = rPr.find(qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"))
        if fill and fill.lower() not in ("auto", "ffffff", "fffffff"):
            # shd fill 为 RRGGBB hex，直接返回 hex 以便动态映射
            return fill.strip().lower()
    return None


def _run_color(run) -> str | None:
    """读取 run 的文本颜色 w:color @w:val，返回 #RRGGBB 或 None（auto）。"""
    rPr = run._element.rPr
    if rPr is None:
        return None
    el = rPr.find(qn("w:color"))
    if el is None:
        return None
    val = el.get(qn("w:val"))
    if not val or val.lower() == "auto":
        return None
    v = val.strip().lstrip("#")
    if len(v) == 6 and all(c in "0123456789abcdefABCDEF" for c in v):
        return f"#{v.lower()}"
    return None


def _run_font_family(run, para=None) -> str | None:
    """读取 run 的字体 w:rFonts（优先 eastAsia > ascii > hAnsi），返回家族名或 None。

    若 run 无直接指定，沿段落直接格式与样式链回退；兼容同一 rPr 内多个 w:rFonts（Word 主题 + 自定义）。
    """
    def _from_rpr(rPr) -> str | None:
        if rPr is None:
            return None
        found = None
        for el in rPr.findall(qn("w:rFonts")):
            for attr in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
                val = el.get(qn(attr))
                if val and val.strip() and val.strip().lower() not in ("theme",):
                    if attr == "w:eastAsia":
                        return val.strip()
                    if found is None:
                        found = val.strip()
            # 兼容 Word 主题字体属性 (majorEastAsia/minorEastAsia)
            for attr in ("w:eastAsiaTheme", "w:asciiTheme", "w:hAnsiTheme", "w:cstheme"):
                theme_val = el.get(qn(attr))
                if theme_val and theme_val.strip():
                    tv = theme_val.strip().lower()
                    if "major" in tv:
                        return "黑体"
                    if "minor" in tv:
                        return "仿宋"
        return found

    # 1. run 直接格式 (w:rPr)
    rPr = getattr(run, "_element", None)
    if rPr is not None:
        rPr = rPr.rPr
    fam = _from_rpr(rPr)
    if fam:
        return fam

    # 2. python-docx run.font.name
    try:
        if run.font and run.font.name and run.font.name.strip():
            return run.font.name.strip()
    except Exception:
        pass

    # 3. 段落直接格式 (w:pPr -> w:rPr)
    if para is not None:
        pPr = getattr(para._p, "pPr", None)
        if pPr is not None:
            p_rpr = pPr.find(qn("w:rPr"))
            fam = _from_rpr(p_rpr)
            if fam:
                return fam

        # 4. 段落样式链
        style = getattr(para, "style", None)
        while style is not None:
            try:
                if style.font and style.font.name and style.font.name.strip():
                    return style.font.name.strip()
            except Exception:
                pass
            sp = style.element.find(qn("w:rPr"))
            fam = _from_rpr(sp)
            if fam:
                return fam
            sp_p = style.element.find(qn("w:pPr"))
            if sp_p is not None:
                fam = _from_rpr(sp_p.find(qn("w:rPr")))
                if fam:
                    return fam
            # 检查内置标题样式名
            sname = (getattr(style, "name", "") or "").strip().lower()
            sid = (getattr(style, "style_id", "") or "").strip().lower()
            if any(sname.startswith(k) or sid.startswith(k) for k in ("heading", "title", "标题", "副标题", "一、", "二、")):
                return "黑体"
            style = getattr(style, "base_style", None)

    return None


def _run_bold(run, para=None, doc=None) -> bool:
    """读取 run 的加粗 w:b / w:bCs（考虑 val=false 显式关闭），沿段落格式与样式链回退。"""
    def _check_b(rPr) -> bool | None:
        if rPr is None:
            return None
        for tag in ("w:b", "w:bCs"):
            for el in rPr.findall(qn(tag)):
                val = el.get(qn("w:val"))
                if val is None:
                    return True  # <w:b/> 无 val 视为 true
                v = val.strip().lower()
                if v in ("false", "0", "off", "none"):
                    return False
                if v in ("true", "1", "on"):
                    return True
                return True
        return None

    # 1. run 直接格式 (w:rPr)
    rPr = getattr(run, "_element", None)
    if rPr is not None:
        rPr = rPr.rPr
    b = _check_b(rPr)
    if b is not None:
        return b

    # 2. python-docx run.bold
    try:
        if run.bold is True or (run.font and run.font.bold is True):
            return True
        if run.bold is False or (run.font and run.font.bold is False):
            return False
    except Exception:
        pass

    # 3. 段落直接格式 (w:pPr -> w:rPr)
    if para is not None:
        pPr = getattr(para._p, "pPr", None)
        if pPr is not None:
            b = _check_b(pPr.find(qn("w:rPr")))
            if b is not None:
                return b

        # 4. 段落样式链
        style = getattr(para, "style", None)
        while style is not None:
            try:
                if style.font and style.font.bold is True:
                    return True
                if style.font and style.font.bold is False:
                    return False
            except Exception:
                pass
            sp = style.element.find(qn("w:rPr"))
            b = _check_b(sp)
            if b is not None:
                return b
            sp_p = style.element.find(qn("w:pPr"))
            if sp_p is not None:
                b = _check_b(sp_p.find(qn("w:rPr")))
                if b is not None:
                    return b
            # 检查内置标题样式名
            sname = (getattr(style, "name", "") or "").strip().lower()
            sid = (getattr(style, "style_id", "") or "").strip().lower()
            if any(sname.startswith(k) or sid.startswith(k) for k in ("heading", "title", "标题", "副标题", "一、", "二、")):
                return True
            style = getattr(style, "base_style", None)

    return False


def _run_font_size_pt(run, para=None, doc=None) -> float | None:
    """读取 run 的字号 w:sz / w:szCs @w:val（half-pt），返回 pt 或 None。

    优先 run 直接格式，其次段落样式链，再次 docDefaults。
    """
    rPr = run._element.rPr
    if rPr is not None:
        for tag in ("w:sz", "w:szCs"):
            last = None
            for el in rPr.findall(qn(tag)):
                val = el.get(qn("w:val"))
                if val and val.strip().isdigit():
                    try:
                        last = int(val.strip()) / 2.0
                    except ValueError:
                        continue
            if last is not None:
                return last
    if para is not None:
        style = para.style
        while style is not None:
            sp = style.element.find(qn("w:rPr"))
            if sp is not None:
                for tag in ("w:sz", "w:szCs"):
                    last = None
                    for el in sp.findall(qn(tag)):
                        val = el.get(qn("w:val"))
                        if val and val.strip().isdigit():
                            try:
                                last = int(val.strip()) / 2.0
                            except ValueError:
                                continue
                    if last is not None:
                        return last
            style = style.base_style
        # 尝试 Normal 样式字体大小
        try:
            sz = style.font.size if style else None  # type: ignore
        except Exception:
            pass
    if doc is not None:
        try:
            # docDefaults
            el = doc.styles.element.find(qn("w:docDefaults"))
            if el is not None:
                el = el.find(qn("w:rPrDefault"))
                if el is not None:
                    el = el.find(qn("w:rPr"))
                    if el is not None:
                        for tag in ("w:sz", "w:szCs"):
                            sz = el.find(qn(tag))
                            if sz is not None:
                                val = sz.get(qn("w:val"))
                                if val and val.strip().isdigit():
                                    return int(val.strip()) / 2.0
        except Exception:
            pass
    return None


def _doc_default_font_pt(doc) -> float:
    """文档默认正文字号（pt）：Normal 样式 > docDefaults > 兜底 12。

    用于把各 Run 的 Word 字号换算为相对比例（正文 ≈ 全局字号）。
    """
    try:
        size = doc.styles["Normal"].font.size
        if size:
            return size.pt
    except KeyError:
        pass
    el = doc.styles.element.find(qn("w:docDefaults"))
    if el is not None:
        el = el.find(qn("w:rPrDefault"))
        if el is not None:
            el = el.find(qn("w:rPr"))
            if el is not None:
                sz = el.find(qn("w:sz"))
                if sz is not None:
                    val = sz.get(qn("w:val"))
                    if val:
                        return int(val) / 2.0  # 半磅 -> pt
    return 12.0


def _split_by_tags(text: str) -> list[tuple[str, str | None]]:
    """按 {{前缀:内容}} 切分文本，返回 [(片段, 标签前缀或None)]。"""
    if not text or "{{" not in text:
        return [(text, None)] if text else []
    res: list[tuple[str, str | None]] = []
    last = 0
    for m in _TAG_RE.finditer(text):
        start, end = m.span()
        if start > last:
            res.append((text[last:start], None))
        key = m.group(1).strip()
        content = m.group(2)
        if content:
            res.append((content, key))
        last = end
    if last < len(text):
        res.append((text[last:], None))
    return [(t, k) for t, k in res if t]


# 固定标签映射（不占用动态 id 池）
_FIXED_TAG_TO_ROLE: dict[str, int] = {
    "手写": 0, "默认手写": 0, "默认": 0,
    "打印": 1, "打印体": 1,
}


def _normalize_tag_key(key: str) -> str:
    return key.strip()


_HEADING_PATTERNS = [
    re.compile(r"^[一二三四五六七八九十百]+、"),          # 一、深化理论学习...
    re.compile(r"^第[一二三四五六七八九十百\d]+[章节目条篇]"),  # 第一章、第一节...
    re.compile(r"^[（\(][一二三四五六七八九十\d]+[）\)]"),    # （一）...
    re.compile(r"^\d+[\.、]\s*\S+"),                     # 1. 1、...
]


def _is_chinese_heading(text: str, align: str = "left", is_first_para: bool = False) -> bool:
    """智能判定中文文档中的标题段落（公文一级/二级标题、首行居中标题等）。"""
    s = text.strip()
    if not s:
        return False
    if len(s) > 60:
        return False
    if align == "center" and len(s) <= 30:
        return True
    if is_first_para and len(s) <= 20:
        return True
    for pat in _HEADING_PATTERNS:
        if pat.match(s):
            return True
    return False


# Word/WPS 标准高亮颜色名称到角色 ID 的精准语义映射（与 GUI _ROLE_BG 底色完全对应）
_HIGHLIGHT_NAME_TO_ROLE_ID: dict[str, int] = {
    "yellow": 2,          # 黄色底 (#fff8b8)
    "darkyellow": 6,      # 橘黄底 (#ffe0b3)
    "green": 3,           # 绿色底 (#d1ffd1)
    "darkgreen": 3,       # 深绿底 (#d1ffd1)
    "cyan": 4,            # 青色/天蓝底 (#c8e8ff)
    "blue": 4,            # 蓝色底 (#c8e8ff)
    "darkblue": 4,        # 深蓝底 (#c8e8ff)
    "darkcyan": 4,        # 深青底 (#c8e8ff)
    "turquoise": 4,       # 绿松石/天蓝 (#c8e8ff)
    "magenta": 5,         # 品红底 (#ffd8f0)
    "pink": 5,            # 粉红底 (#ffd8f0)
    "red": 5,             # 红色底 (#ffd8f0)
    "darkred": 5,         # 深红底 (#ffd8f0)
    "darkmagenta": 7,     # 深品红/紫底 (#e0d8ff)
    "purple": 7,          # 紫色底 (#e0d8ff)
    "violet": 7,          # 紫罗兰底 (#e0d8ff)
    "lightgray": 1,       # 印刷灰 (#e8e8e8)
    "darkgray": 1,        # 印刷灰 (#e8e8e8)
    "gray-25": 1,
    "gray-50": 1,
    "gray": 1,
}


def _hex_to_role_id(hex_str: str) -> int:
    """根据 hex 底纹色彩特征匹配最接近的角色底色（2:黄, 3:绿, 4:蓝, 5:粉, 6:橘, 7:紫）。"""
    s = hex_str.strip().lstrip("#")
    if len(s) != 6:
        return 2
    try:
        r = int(s[0:2], 16)
        g = int(s[2:4], 16)
        b = int(s[4:6], 16)
    except ValueError:
        return 2
    if abs(r - g) < 15 and abs(g - b) < 15:
        return 1 if r < 230 else 0
    if b > r + 30 and b > g - 20:
        return 4
    if g > r + 20 and g > b + 20:
        return 3
    if r > 180 and g > 140 and b < 140:
        return 6 if g < 180 else 2
    if r > b and r > g + 30:
        return 5
    if b > g and r > g:
        return 7
    return 2


_ROLE_DEFAULT_NAMES: dict[int, str] = {
    0: "默认手写",
    1: "打印体",
    2: "手写角色1 (黄)",
    3: "手写角色2 (绿)",
    4: "手写角色3 (蓝)",
    5: "手写角色4 (粉)",
    6: "手写角色5 (橘)",
    7: "手写角色6 (紫)",
}


def has_docx_highlights(path: str | Path) -> bool:
    """快速检查 docx 是否包含任何有效高亮（_run_highlight(run) is not None 且 run.text 非空）。

    包含异常捕获，若读取失败返回 False。
    """
    try:
        doc = Document(str(path))
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text and _run_highlight(run) is not None:
                    return True
        return False
    except Exception:
        return False


def load_paragraphs(
    path: str | Path,
    font_size: int | None = None,
    ignore_highlights: bool = False,
) -> list[Paragraph]:
    """读取 docx 中每个段落，返回 [Paragraph]（忽略空段落）。"""
    paras = load_paragraphs_with_runs(path, font_size=font_size, ignore_highlights=ignore_highlights)
    return paras


def load_paragraphs_with_runs(
    path: str | Path,
    font_size: int | None = None,
    ignore_highlights: bool = False,
) -> list[Paragraph]:
    """读取 docx 返回带 TextRun 的段落列表（支持高亮精准颜色映射与标签）。"""
    doc = Document(str(path))
    # 文档正文基准字号（pt）：Run 的 Word 字号按比例映射到全局字号坐标系，
    # 正文 ≈ 用户设置字号，标题等按文档内比例放大/缩小。
    # 不能换算成 96DPI 绝对像素：高分辨率背景（全局字号 100+）上会小成点。
    body_pt = _doc_default_font_pt(doc) or 12.0
    global_px = float(font_size) if font_size else None
    # 预扫描：文档级是否有任意高亮（决定无标记文本的默认去向）
    has_any_highlight = False
    if not ignore_highlights:
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text and _run_highlight(run) is not None:
                    has_any_highlight = True
                    break
            if has_any_highlight:
                break
    result: list[Paragraph] = []
    highlight_to_role: dict[str, int] = {}
    tag_to_role: dict[str, int] = {}
    used_role_ids: set[int] = {0, 1}
    next_role_id = 2

    def alloc_for_highlight(hl: str) -> int:
        nonlocal next_role_id
        key = hl.lower().strip()
        if key in highlight_to_role:
            return highlight_to_role[key]
        if key in _HIGHLIGHT_NAME_TO_ROLE_ID:
            rid = _HIGHLIGHT_NAME_TO_ROLE_ID[key]
            highlight_to_role[key] = rid
            used_role_ids.add(rid)
            return rid
        clean_hex = key.lstrip("#")
        if len(clean_hex) == 6 and all(c in "0123456789abcdefABCDEF" for c in clean_hex):
            rid = _hex_to_role_id(clean_hex)
            highlight_to_role[key] = rid
            used_role_ids.add(rid)
            return rid
        while next_role_id in used_role_ids:
            next_role_id += 1
        rid = next_role_id
        next_role_id += 1
        highlight_to_role[key] = rid
        used_role_ids.add(rid)
        return rid

    def alloc_for_tag(tag_key: str) -> int:
        nonlocal next_role_id
        nk = _normalize_tag_key(tag_key)
        if nk in _FIXED_TAG_TO_ROLE:
            return _FIXED_TAG_TO_ROLE[nk]
        if nk in tag_to_role:
            return tag_to_role[nk]
        low = nk.lower()
        if low.startswith("角色"):
            suffix = low[2:].strip()
            try:
                n = int(suffix)
                rid = n + 1
                if rid >= next_role_id:
                    next_role_id = rid + 1
                tag_to_role[nk] = rid
                used_role_ids.add(rid)
                return rid
            except ValueError:
                pass
        while next_role_id in used_role_ids:
            next_role_id += 1
        rid = next_role_id
        next_role_id += 1
        tag_to_role[nk] = rid
        used_role_ids.add(rid)
        return rid

    for idx, para in enumerate(doc.paragraphs):
        if not para.text.strip() and not any(r.text for r in para.runs):
            continue
        chars = _first_line_chars(para) or _first_line_emu_chars(para, doc)
        indent = int(round(chars * (font_size or 36))) if chars else 0
        align = _resolve_align(para)
        is_heading_para = _is_chinese_heading(para.text, align=align, is_first_para=(idx == 0))

        if not para.runs:
            plain = para.text.strip()
            if not plain:
                continue
            fallback_role = 1 if has_any_highlight else 0
            fam = "黑体" if is_heading_para else "仿宋"
            ffile = None
            try:
                p = family_to_file(fam)
                if p and Path(p).is_file():
                    ffile = str(p)
            except Exception:
                pass
            result.append(Paragraph(
                text=plain,
                align=align,
                first_line_indent=indent,
                runs=[TextRun(text=plain, role_id=fallback_role, font_family=fam, font_file=ffile, bold=is_heading_para)],
            ))
            continue

        runs: list[TextRun] = []
        for run in para.runs:
            raw = run.text
            if not raw:
                continue
            hl = None if ignore_highlights else _run_highlight(run)
            col = _run_color(run)
            fam = _run_font_family(run, para)
            is_bold = _run_bold(run, para, doc) or is_heading_para
            if is_bold or is_heading_para:
                if not fam or fam.strip().lower() in ("仿宋", "仿宋_gb2312", "fangsong", "simfang", "宋体", "simsun"):
                    fam = "黑体"
            elif not fam:
                fam = "仿宋"
            pt = _run_font_size_pt(run, para, doc)
            if pt is not None and global_px:
                fpx = max(1, int(round(pt / body_pt * global_px)))
            else:
                fpx = None
            ffile = None
            if fam:
                try:
                    p = family_to_file(fam)
                    if p and Path(p).is_file():
                        ffile = str(p)
                except Exception:
                    ffile = None
            segments = _split_by_tags(raw)
            if not segments:
                continue
            def _font_kwargs_for(role_id: int) -> dict:
                if role_id == 1:
                    return dict(font_family=fam, font_size=fpx, font_file=ffile, bold=is_bold)
                return dict(font_family=None, font_size=None, font_file=None, bold=False)

            if len(segments) == 1 and segments[0][1] is None:
                text, _ = segments[0]
                if hl is None:
                    role = 1 if has_any_highlight else 0
                else:
                    role = alloc_for_highlight(hl)
                runs.append(TextRun(text=text, role_id=role, color=col, **_font_kwargs_for(role)))
            else:
                for seg_text, seg_key in segments:
                    if seg_key is None:
                        if hl is None:
                            role = 1 if has_any_highlight else 0
                        else:
                            role = alloc_for_highlight(hl)
                        runs.append(TextRun(text=seg_text, role_id=role, color=col, **_font_kwargs_for(role)))
                    else:
                        # 全部手写模式下忽略标签的角色指令，仅保留剥离语法后的文本
                        role = 0 if ignore_highlights else alloc_for_tag(seg_key)
                        runs.append(TextRun(text=seg_text, role_id=role, color=col, **_font_kwargs_for(role)))
        if not runs:
            continue
        merged: list[TextRun] = []
        for r in runs:
            if merged and merged[-1].role_id == r.role_id and merged[-1].color == r.color \
               and merged[-1].font_family == r.font_family and merged[-1].font_size == r.font_size \
               and merged[-1].font_file == r.font_file and merged[-1].bold == r.bold:
                merged[-1].text += r.text
            else:
                merged.append(r)
        plain = "".join(r.text for r in merged)
        if not plain.strip():
            continue
        result.append(Paragraph(text=plain, align=align, first_line_indent=indent, runs=merged))
    return result


def extract_roles_from_paragraphs(paragraphs: list[Paragraph]) -> list[HandwritingRole]:
    """从已解析的段落 runs 中推断角色表（供 GUI 初始化角色面板）。

    返回动态分配的角色列表（包含 0/1 固定角色）。
    """
    seen: dict[int, str] = {0: "默认手写", 1: "打印体"}
    for p in paragraphs:
        for r in p.runs or []:
            if r.role_id not in seen:
                seen[r.role_id] = _ROLE_DEFAULT_NAMES.get(r.role_id, f"手写角色{r.role_id-1}")
    roles = []
    for rid in sorted(seen):
        roles.append(HandwritingRole(id=rid, name=seen[rid], printed=(rid == 1)))
    return roles
