"""测试 docx 标题加粗识别与中文字体推断（黑体/仿宋/自定义）。"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from handwritesim.core.docx_io import load_paragraphs_with_runs, _run_bold, _run_font_family
from handwritesim.core.system_fonts import family_to_file


def _set_cell_rfonts(rPr, eastAsia=None, eastAsiaTheme=None):
    rFonts = OxmlElement("w:rFonts")
    if eastAsia:
        rFonts.set(qn("w:eastAsia"), eastAsia)
    if eastAsiaTheme:
        rFonts.set(qn("w:eastAsiaTheme"), eastAsiaTheme)
    rPr.append(rFonts)


def test_docx_bold_and_fonts_heading_and_body():
    """测试带有加粗标题与正文的 docx 解析：标题为黑体加粗，正文为仿宋。"""
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "test_doc.docx"
        doc = Document()

        # 段落 1：标题（居中，加粗，无背景）
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run("思想汇报")
        r1.bold = True

        # 段落 2：正文（首行缩进，未加粗）
        p2 = doc.add_paragraph()
        r2 = p2.add_run("敬爱的党组织：")

        # 段落 3：带高亮的手写段落（标记）
        p3 = doc.add_paragraph()
        r3_print = p3.add_run("本季度我围绕")
        r3_hw = p3.add_run("深入学习领会")
        # 添加高亮黄色
        rPr_hw = r3_hw._element.get_or_add_rPr()
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), "yellow")
        rPr_hw.append(hl)

        doc.save(str(p))

        paras = load_paragraphs_with_runs(p, font_size=36)
        assert len(paras) == 3

        # 段落 1 标题：打印体，加粗，黑体
        runs1 = paras[0].runs
        assert runs1 is not None and len(runs1) == 1
        assert runs1[0].role_id == 1  # 存在高亮，无标记段为打印体
        assert runs1[0].bold is True
        assert runs1[0].font_family == "黑体"
        assert runs1[0].font_file is not None and "simhei" in runs1[0].font_file.lower()

        # 段落 2 正文：打印体，未加粗，仿宋
        runs2 = paras[1].runs
        assert runs2 is not None and len(runs2) == 1
        assert runs2[0].role_id == 1
        assert runs2[0].bold is False
        assert runs2[0].font_family == "仿宋"
        assert runs2[0].font_file is not None and "simfang" in runs2[0].font_file.lower()

        # 段落 3：打印体 + 手写角色 2
        runs3 = paras[2].runs
        assert runs3 is not None and len(runs3) == 2
        assert runs3[0].role_id == 1  # 打印体
        assert runs3[0].bold is False
        assert runs3[1].role_id == 2  # 手写角色


def test_docx_paragraph_level_bold():
    """测试段落级 pPr/rPr 直接加粗。"""
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "test_ppr_bold.docx"
        doc = Document()
        p1 = doc.add_paragraph()
        # 在 pPr 中注入 rPr -> w:b
        pPr = p1._p.get_or_add_pPr()
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        pPr.append(rPr)
        r1 = p1.add_run("段落直接加粗的标题")

        # 增加一个黄色高亮使得文档进入打印/手写混排模式
        p2 = doc.add_paragraph()
        r2 = p2.add_run("手写部分")
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), "yellow")
        r2._element.get_or_add_rPr().append(hl)

        doc.save(str(p))

        paras = load_paragraphs_with_runs(p, font_size=36)
        runs1 = paras[0].runs
        assert runs1 is not None and len(runs1) == 1
        assert runs1[0].bold is True
        assert runs1[0].font_family == "黑体"


def test_docx_run_font_size_relative_to_global_setting():
    """Run 字号应映射到全局字号坐标系：正文 1:1，标题按文档比例放大。

    此前实现把 Word 字号按 96DPI 换算成绝对像素（10.5pt→14px），
    在高分辨率背景（4200px 宽、全局字号 100+）上打印文字会小成点。
    """
    with tempfile.TemporaryDirectory() as td:
        from docx.shared import Pt

        p = Path(td) / "test_font_size.docx"
        doc = Document()
        doc.styles["Normal"].font.size = Pt(10.5)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = title.add_run("思想汇报")
        rt.bold = True
        rt.font.size = Pt(22)

        body = doc.add_paragraph()
        body.add_run("敬爱的党组织：")

        # 高亮使文档进入打印/手写混排模式
        hw = doc.add_paragraph()
        r_hw = hw.add_run("手写段")
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), "yellow")
        r_hw._element.get_or_add_rPr().append(hl)

        doc.save(str(p))

        paras = load_paragraphs_with_runs(p, font_size=36)
        # 标题 22pt：22 / 10.5 * 36 ≈ 75
        assert paras[0].runs[0].font_size == 75
        # 正文（无显式字号，随文档默认 10.5pt）：与全局字号 1:1
        assert paras[1].runs[0].font_size == 36

        # 全局字号变化时按比例缩放
        paras2 = load_paragraphs_with_runs(p, font_size=100)
        assert paras2[0].runs[0].font_size == 210
        assert paras2[1].runs[0].font_size == 100



    """测试 majorEastAsia / minorEastAsia 主题字体识别。"""
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "test_theme_fonts.docx"
        doc = Document()
        p1 = doc.add_paragraph()
        r1 = p1.add_run("大标题")
        rPr1 = r1._element.get_or_add_rPr()
        _set_cell_rfonts(rPr1, eastAsiaTheme="majorEastAsia")

        p2 = doc.add_paragraph()
        r2 = p2.add_run("正文文字")
        rPr2 = r2._element.get_or_add_rPr()
        _set_cell_rfonts(rPr2, eastAsiaTheme="minorEastAsia")

        # 高亮
        p3 = doc.add_paragraph()
        r3 = p3.add_run("手写")
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), "cyan")
        r3._element.get_or_add_rPr().append(hl)

        doc.save(str(p))

        paras = load_paragraphs_with_runs(p, font_size=36)
        assert paras[0].runs[0].font_family == "黑体"
        assert paras[1].runs[0].font_family == "仿宋"


def test_docx_chinese_heading_patterns_bold_and_simhei():
    """测试中文公文标题结构（思想汇报、一、二、三、）自动识别加粗与黑体。"""
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "test_headings.docx"
        doc = Document()

        p0 = doc.add_paragraph()
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.add_run("思想汇报")

        p1 = doc.add_paragraph()
        p1.add_run("敬爱的党组织：")

        p2 = doc.add_paragraph()
        p2.add_run("2025年第三季度，恰逢中国共产党成立104周年...")

        p3 = doc.add_paragraph()
        p3.add_run("一、深化理论学习，筑牢思想根基")

        p4 = doc.add_paragraph()
        r4_1 = p4.add_run("本季度我围绕")
        r4_2 = p4.add_run("“七一”和抗战胜利80周年")
        hl_y = OxmlElement("w:highlight")
        hl_y.set(qn("w:val"), "yellow")
        r4_2._element.get_or_add_rPr().append(hl_y)

        p5 = doc.add_paragraph()
        p5.add_run("二、立足本职岗位，践行党员担当")

        p6 = doc.add_paragraph()
        p6.add_run("三、坚持自我剖析，正视自身不足")

        p7 = doc.add_paragraph()
        r7_1 = p7.add_run("对照党员标准，我仍存在不足：")
        r7_2 = p7.add_run("一是党史学习的深度还不够")
        hl_b = OxmlElement("w:highlight")
        hl_b.set(qn("w:val"), "blue")
        r7_2._element.get_or_add_rPr().append(hl_b)

        doc.save(str(p))

        paras = load_paragraphs_with_runs(p, font_size=36)
        # 思想汇报 (首行居中标题) -> 黑体加粗
        assert paras[0].runs[0].bold is True
        assert paras[0].runs[0].font_family == "黑体"
        assert paras[0].runs[0].role_id == 1

        # 敬爱的党组织： -> 普通正文
        assert paras[1].runs[0].bold is False
        assert paras[1].runs[0].font_family == "仿宋"

        # 一、深化理论学习... -> 黑体加粗
        assert paras[3].runs[0].bold is True
        assert paras[3].runs[0].font_family == "黑体"

        # 黄色高亮 -> role 2 (黄)
        assert paras[4].runs[1].role_id == 2

        # 二、立足本职岗位... -> 黑体加粗
        assert paras[5].runs[0].bold is True
        assert paras[5].runs[0].font_family == "黑体"

        # 三、坚持自我剖析... -> 黑体加粗
        assert paras[6].runs[0].bold is True
        assert paras[6].runs[0].font_family == "黑体"

        # 蓝色高亮 -> role 4 (蓝) 而不是 3 (绿)
        assert paras[7].runs[1].role_id == 4

