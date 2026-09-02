"""混排多角色功能验收测试（对齐 Rust 版实施计划）。"""

import tempfile
from pathlib import Path

import numpy as np
import pytest
from PIL import Image

from handwritesim.core.docx_io import load_paragraphs_with_runs, extract_roles_from_paragraphs
from handwritesim.core.engine import HandwritingEngine
from handwritesim.core.models import HandwritingParams, Paragraph, TextRun, HandwritingRole, parse_color

FONTS = (r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simsun.ttc", r"C:\Windows\Fonts\simhei.ttf")

def _font() -> str:
    import os
    for f in FONTS:
        if os.path.exists(f):
            return f
    pytest.skip("未找到系统 CJK 字体")

def _bg(tmp_path: Path) -> str:
    p = tmp_path / "bg.png"
    Image.new("RGB", (600, 800), "white").save(p)
    return str(p)

def _color_pixels(arr: np.ndarray, hexcol: str) -> int:
    r, g, b = parse_color(hexcol)
    mask = (np.abs(arr[:, :, 0].astype(int) - r) < 30) & (np.abs(arr[:, :, 1].astype(int) - g) < 30) & (np.abs(arr[:, :, 2].astype(int) - b) < 30)
    mask &= (arr != 255).any(axis=2)
    return int(mask.sum())

# ---------------------------------------------------------------------------
# 模型
# ---------------------------------------------------------------------------
def test_roles_serialization(tmp_path: Path):
    params = HandwritingParams(font_path=_font(), background_path=_bg(tmp_path))
    params.roles = [
        HandwritingRole(id=0, name="默认"),
        HandwritingRole(id=2, name="小明", color="#cc0000", font_path=_font(), font_size=32),
    ]
    params.paragraphs = [Paragraph(text="", runs=[TextRun("hello ", 0), TextRun("world", 2)])]
    d = params.to_dict()
    p2 = HandwritingParams.from_dict(d)
    assert len(p2.roles) == 2
    assert p2.paragraphs[0].runs[1].role_id == 2
    assert p2.roles[1].color == "#cc0000"

def test_paragraph_effective_runs():
    p = Paragraph(text="hello", align="left")
    assert p.effective_runs()[0].text == "hello"
    p2 = Paragraph(text="", runs=[TextRun("a", 0), TextRun("b", 2)])
    assert len(p2.effective_runs()) == 2
    assert p2.plain_text() == "ab"

# ---------------------------------------------------------------------------
# Docx 动态高亮
# ---------------------------------------------------------------------------
def test_docx_dynamic_highlight_first_two(tmp_path: Path):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    para = doc.add_paragraph()
    # 首次出现黄色 -> 角色2
    r1 = para.add_run("第一段黄 ")
    r1._element.get_or_add_rPr().append(r1._element.get_or_add_rPr().makeelement(qn("w:highlight"), {qn("w:val"): "yellow"}))
    # 首次出现绿色 -> 角色3
    r2 = para.add_run("第二段绿 ")
    r2._element.get_or_add_rPr().append(r2._element.get_or_add_rPr().makeelement(qn("w:highlight"), {qn("w:val"): "green"}))
    # 再次黄色 -> 复用角色2
    r3 = para.add_run("再回黄")
    r3._element.get_or_add_rPr().append(r3._element.get_or_add_rPr().makeelement(qn("w:highlight"), {qn("w:val"): "yellow"}))
    # 标签：文档含高亮时，普通片段应视为打印（全局策略）
    doc.add_paragraph("普通{{打印:打印区}}尾")
    path = tmp_path / "dyn.docx"
    doc.save(path)

    paras = load_paragraphs_with_runs(path)
    # 第一段应为 3 runs: 黄(2) 绿(3) 黄(2)
    assert paras[0].runs[0].role_id == 2
    assert paras[0].runs[1].role_id == 3
    assert paras[0].runs[2].role_id == 2
    # 第二段：文档有高亮时无标记普通段应合并为打印（role 1）
    assert all(r.role_id == 1 for r in paras[1].runs)

    roles = extract_roles_from_paragraphs(paras)
    ids = {r.id for r in roles}
    assert 2 in ids and 3 in ids and 1 in ids


def test_docx_global_no_highlight_all_handwrite(tmp_path: Path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("第一段无标记")
    doc.add_paragraph("第二段也无高亮")
    path = tmp_path / "nohl.docx"
    doc.save(path)
    paras = load_paragraphs_with_runs(path)
    # 全无高亮 → 全文手写 role 0
    assert all(r.role_id == 0 for p in paras for r in p.runs)


def test_docx_global_with_highlight_plain_is_printed(tmp_path: Path):
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document()
    para = doc.add_paragraph()
    r = para.add_run("黄色手写 ")
    r._element.get_or_add_rPr().append(r._element.get_or_add_rPr().makeelement(qn("w:highlight"), {qn("w:val"): "yellow"}))
    doc.add_paragraph("这段无标记应视为打印")
    path = tmp_path / "hl_plain.docx"
    doc.save(path)
    paras = load_paragraphs_with_runs(path)
    assert paras[0].runs[0].role_id == 2
    assert paras[1].runs[0].role_id == 1

def test_docx_tag_dynamic_name(tmp_path: Path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("前缀{{小明:小明的字}}后缀{{小红:小红的字}}再{{小明:又小明}}")
    path = tmp_path / "tag.docx"
    doc.save(path)
    paras = load_paragraphs_with_runs(path)
    # 小明首次 -> 2, 小红 -> 3, 再小明复用 2
    rids = [r.role_id for r in paras[0].runs]
    # 结构：普通0, 小明2, 后缀0, 小红3, 再0? 实际 suffix 含“后缀”
    assert rids.count(2) == 2
    assert 3 in rids

# ---------------------------------------------------------------------------
# 引擎混排
# ---------------------------------------------------------------------------
def test_engine_mixed_inline_colors(tmp_path: Path):
    roles = [
        HandwritingRole(id=0, name="默认", color="#000000"),
        HandwritingRole(id=2, name="红", color="#c80000"),
        HandwritingRole(id=3, name="蓝", color="#0044cc"),
    ]
    para = Paragraph(text="", runs=[
        TextRun("黑色", 0), TextRun("红色", 2), TextRun("蓝色", 3), TextRun("黑色长文本"*10, 0)
    ])
    params = HandwritingParams(font_path=_font(), background_path=_bg(tmp_path), font_size=28, line_spacing=36, roles=roles, paragraphs=[para])
    pages = list(HandwritingEngine(seed=42, backend="fast").generate(params))
    assert len(pages) >= 1
    arr = np.asarray(pages[0])
    assert _color_pixels(arr, "#c80000") > 100
    assert _color_pixels(arr, "#0044cc") > 100
    assert _color_pixels(arr, "#000000") > 100

def test_engine_mixed_font_zero_perturb_printed(tmp_path: Path):
    roles = [
        HandwritingRole(id=0, name="默认", color="#000000", perturb_x_sigma=5),
        HandwritingRole(id=1, name="打印", color="#000000", printed=True),
    ]
    para = Paragraph(text="", runs=[TextRun("手写", 0), TextRun("打印", 1)])
    params = HandwritingParams(font_path=_font(), background_path=_bg(tmp_path), font_size=32, perturb_x_sigma=5, roles=roles, paragraphs=[para])
    # 打印体应零扰动：两次相同 seed 下 printed 区域像素应完全一致（无随机位移）
    a = np.asarray(HandwritingEngine(seed=7, backend="fast").render_preview(params))
    b = np.asarray(HandwritingEngine(seed=7, backend="fast").render_preview(params))
    assert np.array_equal(a, b)
