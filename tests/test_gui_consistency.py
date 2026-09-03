"""GUI 预览/导出一致性回归测试（依赖 Qt 会话）。"""

from __future__ import annotations

import os
from pathlib import Path

import pytest
from PIL import Image
from PyQt6.QtWidgets import QApplication, QFileDialog, QMessageBox

from handwritesim.gui.main_window import MainWindow

_FONTS = (
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\simsun.ttc",
)


def _font() -> str:
    for font in _FONTS:
        if os.path.exists(font):
            return font
    pytest.skip("未找到系统 CJK 字体")


@pytest.fixture(scope="module")
def qapp():
    app = QApplication.instance() or QApplication([])
    return app


@pytest.fixture()
def window(qapp, tmp_path: Path) -> MainWindow:
    win = MainWindow(out_dir=tmp_path / "out")
    win._ui.lineEdit.setText(_font())
    bg = tmp_path / "bg.png"
    Image.new("RGB", (600, 400), "white").save(bg)
    win._ui.lineEdit_2.setText(str(bg))
    win._ui.textEdit.setPlainText("预览导出一致性测试。")
    return win


class _BusyWorker:
    """模拟正在运行的 RenderWorker。"""

    def isRunning(self) -> bool:
        return True


def test_export_uses_last_preview_snapshot(window, monkeypatch) -> None:
    """导出应复用最后一次预览的参数与种子，而非重新收集界面参数。

    回归：竞态下（预览渲染期间用户又输入），自动预览被跳过时屏幕仍是
    旧内容，但旧实现更新了 seed 且导出重新收集参数，导致导出与预览
    不一致（长文本时第二页差一个字符）。
    """
    old_params = window.collect_params()
    window._preview_seed = 777
    window._preview_params = old_params

    # 用户又输入了内容（预览被跳过，屏幕仍显示旧内容）
    window._ui.textEdit.setPlainText("预览导出一致性测试。补充文字。")

    captured = {}
    monkeypatch.setattr(
        window,
        "_start_worker",
        lambda params, mode, quiet=False, seed=None: captured.update(
            params=params, seed=seed
        )
        or True,
    )
    window._on_export()

    assert captured["seed"] == 777
    assert captured["params"] is old_params  # 预览时的参数快照


def test_skipped_auto_preview_keeps_snapshot(window, monkeypatch) -> None:
    """自动预览因 worker 忙被跳过时，不得更新 seed/参数快照。"""
    old_params = window.collect_params()
    window._preview_seed = 555
    window._preview_params = old_params
    window._ui.textEdit.setPlainText("被跳过的预览文本。")

    monkeypatch.setattr(window, "_worker", _BusyWorker())
    monkeypatch.setattr(window, "_set_busy", lambda busy: None)
    window._on_auto_preview()

    assert window._preview_seed == 555
    assert window._preview_params is old_params


def test_preview_success_updates_snapshot(window, monkeypatch) -> None:
    """预览真正启动后应更新种子与参数快照。"""
    monkeypatch.setattr(window, "_worker", None)
    monkeypatch.setattr(window, "_set_busy", lambda busy: None)

    def fake_start_worker(params, mode, quiet=False, seed=None):
        window._worker = object()
        return True

    monkeypatch.setattr(window, "_start_worker", fake_start_worker)
    window._on_auto_preview()

    assert window._preview_seed is not None
    assert window._preview_params is not None
    assert window._preview_params.text == "预览导出一致性测试。"


def test_export_without_preview_uses_current_params(window, monkeypatch) -> None:
    """从未预览过时，导出应使用当前界面参数且 seed 为 None（保持随机）。"""
    assert window._preview_params is None
    captured = {}
    monkeypatch.setattr(
        window,
        "_start_worker",
        lambda params, mode, quiet=False, seed=None: captured.update(
            params=params, seed=seed
        )
        or True,
    )
    window._on_export()

    assert captured["seed"] is None
    assert captured["params"].text == "预览导出一致性测试。"


def test_preset_combo_wheel_does_not_switch(qapp, tmp_path: Path) -> None:
    """预设下拉框必须点击切换：滚轮滚动不得触发切换（避免误触）。

    预设下拉框位于可滚动面板内，若滚轮直接切换预设，用户滚动参数面板时
    会误触发配置变更，故需忽略滚轮事件（与 NoWheelSpinBox 同策略）。
    """
    from PyQt6.QtCore import QPoint, QPointF, Qt
    from PyQt6.QtGui import QWheelEvent

    from handwritesim.gui.ui import NoWheelComboBox

    win = MainWindow(out_dir=tmp_path / "out")
    combo = win._ui.combo_preset
    assert isinstance(combo, NoWheelComboBox)

    # 构造向下滚动事件（angleDelta.y > 0）：默认行为会切换选项
    event = QWheelEvent(
        QPointF(0, 0),
        QPointF(0, 0),
        QPoint(0, 0),
        QPoint(0, 120),
        Qt.MouseButton.NoButton,
        Qt.KeyboardModifier.NoModifier,
        Qt.ScrollPhase.NoScrollPhase,
        False,
    )
    combo.setCurrentIndex(0)
    combo.wheelEvent(event)
    assert combo.currentIndex() == 0  # 滚轮不得改变当前项


def test_region_dialog_features(qapp) -> None:
    """测试 RegionDialog 的参数回填、段落工具栏与高级覆盖面板。"""
    from handwritesim.core.models import Paragraph
    from handwritesim.gui.region_dialog import RegionDialog

    dlg = RegionDialog(
        title="编辑文字区域",
        text="段落一\n段落二",
        paragraphs=[
            Paragraph(text="段落一", align="center", first_line_indent=40),
            Paragraph(text="段落二", align="right", first_line_indent=0),
        ],
        printed=True,
        font_path="dummy.ttf",
        font_size=28,
        page=2,
        word_spacing=10,
        line_spacing=50,
        perturb_theta_sigma=0.08,
        miswrite_rate=0.15,
        miswrite_strikeout_style="slash",
        color="#ff0000",
        margin_top=12,
        margin_left=15,
    )

    assert dlg.region_printed is True
    assert dlg.region_font_size == 28
    assert dlg.region_page == 2
    assert dlg.region_word_spacing == 10
    assert dlg.region_line_spacing == 50
    assert dlg.region_perturb_theta_sigma == 0.08
    assert dlg.region_miswrite_rate == 0.15
    assert dlg.region_miswrite_strikeout_style == "slash"
    assert dlg.region_color == "#ff0000"
    assert dlg.region_margin_top == 12
    assert dlg.region_margin_left == 15

    paras = dlg.region_paragraphs
    assert len(paras) == 2
    assert paras[0].text == "段落一"
    assert paras[0].align == "center"
    assert paras[1].text == "段落二"
    assert paras[1].align == "right"


def test_downscaled_preview_params_preserves_regions_and_paragraphs(window) -> None:
    """测试预览参数降采样时，完整保留区域与主文本的段落排版、对齐及覆盖项。"""
    from handwritesim.core.models import HandwritingParams, Paragraph, TextRegion

    bg_path = window._ui.lineEdit_2.text()
    p = HandwritingParams(
        font_path=window._ui.lineEdit.text(),
        background_path=bg_path,
        font_size=40,
        paragraphs=[
            Paragraph(text="主段落", align="center", first_line_indent=20)
        ],
        regions=[
            TextRegion(
                x=100, y=100, w=200, h=100,
                text="你好\n中刺刀",
                align="left",
                paragraphs=[
                    Paragraph(text="你好", align="left"),
                    Paragraph(text="中刺刀", align="center"),
                ],
                color="#123456",
                margin_left=10,
                miswrite_strikeout_style="cross",
            )
        ]
    )
    # 模拟图片宽度很大触发降采样
    window._preview_max_width = 300
    preview = window._downsample_preview(p)

    assert preview.paragraphs is not None
    assert preview.paragraphs[0].align == "center"

    assert preview.regions is not None and len(preview.regions) == 1
    r = preview.regions[0]
    assert r.color == "#123456"
    assert r.miswrite_strikeout_style == "cross"
    assert r.paragraphs is not None and len(r.paragraphs) == 2
    assert r.paragraphs[0].align == "left"
    assert r.paragraphs[1].text == "中刺刀"
    assert r.paragraphs[1].align == "center"


def _make_highlighted_docx(path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Normal text ")
    r2 = p.add_run("Highlighted text")
    r2.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.save(str(path))


def _make_plain_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Plain text only")
    doc.save(str(path))


def test_import_docx_all_handwriting(window, monkeypatch, tmp_path: Path) -> None:
    """带高亮 docx 导入正文：选择‘全部作为手写’，所有 runs 为 role_id 0 且不生成额外角色。"""
    docx_path = tmp_path / "highlighted.docx"
    _make_highlighted_docx(docx_path)

    monkeypatch.setattr(
        QFileDialog,
        "getOpenFileName",
        lambda *args, **kwargs: (str(docx_path), "Word 文档 (*.docx)"),
    )

    box_instances = []

    def fake_exec(self):
        box_instances.append(self)
        for btn in self.buttons():
            if "全部作为手写" in btn.text():
                self._clicked_btn = btn
                break
        return 0

    monkeypatch.setattr(QMessageBox, "exec", fake_exec)
    monkeypatch.setattr(
        QMessageBox, "clickedButton", lambda self: getattr(self, "_clicked_btn", None)
    )

    window._import_docx()

    assert len(box_instances) == 1
    box = box_instances[0]
    assert box.windowTitle() == "检测到文字高亮标记"
    assert "全部作为手写（推荐）" in [b.text() for b in box.buttons()]

    params = window.collect_params()
    assert params.paragraphs is not None and len(params.paragraphs) > 0
    assert "Highlighted text" in window._ui.textEdit.toPlainText()
    # 验证全部段落均为默认手写（runs 为 None 或 runs 全部为 role_id 0）
    for p in params.paragraphs:
        if p.runs:
            assert all(r.role_id == 0 for r in p.runs)
    # 验证编辑器中的字符格式属性均非打印体/高亮角色 (role_id 0 或 None)
    from PyQt6.QtGui import QTextCharFormat
    doc = window._ui.textEdit.document()
    for i in range(doc.blockCount()):
        it = doc.findBlockByNumber(i).begin()
        while not it.atEnd():
            frag = it.fragment()
            if frag.isValid():
                cf = frag.charFormat()
                rid_val = cf.property(QTextCharFormat.Property.UserProperty)
                assert rid_val in (None, 0)
            it += 1
    # 角色列表中不包含高亮动态角色 (role_id >= 2)
    assert not any(r.id >= 2 for r in window._roles)


def test_import_docx_mixed(window, monkeypatch, tmp_path: Path) -> None:
    """带高亮 docx 导入正文：选择‘打印/手写混排’，识别打印体与高亮笔迹角色。"""
    docx_path = tmp_path / "highlighted.docx"
    _make_highlighted_docx(docx_path)

    monkeypatch.setattr(
        QFileDialog,
        "getOpenFileName",
        lambda *args, **kwargs: (str(docx_path), "Word 文档 (*.docx)"),
    )

    box_instances = []

    def fake_exec(self):
        box_instances.append(self)
        for btn in self.buttons():
            if "打印/手写混排" in btn.text():
                self._clicked_btn = btn
                break
        return 0

    monkeypatch.setattr(QMessageBox, "exec", fake_exec)
    monkeypatch.setattr(
        QMessageBox, "clickedButton", lambda self: getattr(self, "_clicked_btn", None)
    )

    window._import_docx()

    assert len(box_instances) == 1
    assert box_instances[0].windowTitle() == "检测到文字高亮标记"

    params = window.collect_params()
    assert params.paragraphs is not None and len(params.paragraphs) > 0
    all_role_ids = {
        r.role_id for p in params.paragraphs for r in (p.runs or [])
    }
    assert 1 in all_role_ids  # 未标记文本为打印体 (role_id 1)
    assert any(rid >= 2 for rid in all_role_ids)  # 高亮文本生成高亮角色 (role_id >= 2)


def test_import_docx_cancel(window, monkeypatch, tmp_path: Path) -> None:
    """带高亮 docx 导入正文：点击‘取消’，退出导入且不修改编辑器内容。"""
    window._ui.textEdit.setPlainText("原始正文保留不被覆盖。")
    docx_path = tmp_path / "highlighted.docx"
    _make_highlighted_docx(docx_path)

    monkeypatch.setattr(
        QFileDialog,
        "getOpenFileName",
        lambda *args, **kwargs: (str(docx_path), "Word 文档 (*.docx)"),
    )

    box_instances = []

    def fake_exec(self):
        box_instances.append(self)
        for btn in self.buttons():
            if "取消" in btn.text():
                self._clicked_btn = btn
                break
        return 0

    monkeypatch.setattr(QMessageBox, "exec", fake_exec)
    monkeypatch.setattr(
        QMessageBox, "clickedButton", lambda self: getattr(self, "_clicked_btn", None)
    )

    window._import_docx()

    assert len(box_instances) == 1
    assert box_instances[0].windowTitle() == "检测到文字高亮标记"
    assert window._ui.textEdit.toPlainText() == "原始正文保留不被覆盖。"


def test_import_docx_without_highlights_no_dialog(
    window, monkeypatch, tmp_path: Path
) -> None:
    """无高亮 docx 导入正文：静默导入，不弹确认对话框。"""
    docx_path = tmp_path / "plain.docx"
    _make_plain_docx(docx_path)

    monkeypatch.setattr(
        QFileDialog,
        "getOpenFileName",
        lambda *args, **kwargs: (str(docx_path), "Word 文档 (*.docx)"),
    )

    box_instances = []

    def fake_exec(self):
        box_instances.append(self)
        return 0

    monkeypatch.setattr(QMessageBox, "exec", fake_exec)
    monkeypatch.setattr(QMessageBox, "clickedButton", lambda self: None)

    window._import_docx()

    highlight_dialogs = [
        b for b in box_instances if b.windowTitle() == "检测到文字高亮标记"
    ]
    assert len(highlight_dialogs) == 0
    assert "Plain text only" in window._ui.textEdit.toPlainText()


def test_import_document_keep_raw(window, monkeypatch, tmp_path: Path) -> None:
    """底图导入：检测到标记时选择‘保留完整底图’，底图指向 _raw 且区域列表为空。"""
    from handwritesim.core.models import TextRegion
    import handwritesim.core.doc_render as doc_render

    doc_file = tmp_path / "mock.pdf"
    doc_file.touch()
    page_img = tmp_path / "page_0.png"
    raw_img = tmp_path / "page_0_raw.png"
    page_img.touch()
    raw_img.touch()

    region = TextRegion(
        x=10, y=10, w=100, h=40, page=1, role_id=2, highlight="yellow"
    )

    monkeypatch.setattr(
        QFileDialog,
        "getOpenFileName",
        lambda *args, **kwargs: (str(doc_file), "文档 (*.pdf *.docx)"),
    )
    monkeypatch.setattr(
        QMessageBox, "question", lambda *args, **kwargs: QMessageBox.StandardButton.Yes
    )
    monkeypatch.setattr(QMessageBox, "information", lambda *args, **kwargs: None)
    monkeypatch.setattr(QMessageBox, "warning", lambda *args, **kwargs: None)
    monkeypatch.setattr(
        doc_render,
        "document_to_page_images_with_regions",
        lambda path, out_dir, dpi=200: ([page_img], [region]),
    )

    box_instances = []

    def fake_exec(self):
        box_instances.append(self)
        for btn in self.buttons():
            if "保留完整底图" in btn.text():
                self._clicked_btn = btn
                break
        return 0

    monkeypatch.setattr(QMessageBox, "exec", fake_exec)
    monkeypatch.setattr(
        QMessageBox, "clickedButton", lambda self: getattr(self, "_clicked_btn", None)
    )

    window._import_document()

    assert len(box_instances) == 1
    box = box_instances[0]
    assert box.windowTitle() == "检测到手写填空标记"
    assert "保留完整底图" in [b.text() for b in box.buttons()]

    assert window._doc_pages == [str(raw_img)]
    assert window._regions == []
    assert "完整底图" in window._ui.lineEdit_14.text()
    assert window._ui.lineEdit_2.text() == str(raw_img)


def test_import_document_extract_regions(window, monkeypatch, tmp_path: Path) -> None:
    """底图导入：检测到标记时选择‘提取填空框’，提取区域并同步角色。"""
    from handwritesim.core.models import TextRegion
    import handwritesim.core.doc_render as doc_render

    doc_file = tmp_path / "mock.pdf"
    doc_file.touch()
    page_img = tmp_path / "page_0.png"
    raw_img = tmp_path / "page_0_raw.png"
    page_img.touch()
    raw_img.touch()

    region = TextRegion(
        x=10, y=10, w=100, h=40, page=1, role_id=2, highlight="yellow"
    )

    monkeypatch.setattr(
        QFileDialog,
        "getOpenFileName",
        lambda *args, **kwargs: (str(doc_file), "文档 (*.pdf *.docx)"),
    )
    monkeypatch.setattr(
        QMessageBox, "question", lambda *args, **kwargs: QMessageBox.StandardButton.Yes
    )
    monkeypatch.setattr(QMessageBox, "information", lambda *args, **kwargs: None)
    monkeypatch.setattr(
        doc_render,
        "document_to_page_images_with_regions",
        lambda path, out_dir, dpi=200: ([page_img], [region]),
    )

    box_instances = []

    def fake_exec(self):
        box_instances.append(self)
        for btn in self.buttons():
            if "提取填空框" in btn.text():
                self._clicked_btn = btn
                break
        return 0

    monkeypatch.setattr(QMessageBox, "exec", fake_exec)
    monkeypatch.setattr(
        QMessageBox, "clickedButton", lambda self: getattr(self, "_clicked_btn", None)
    )

    window._import_document()

    assert len(box_instances) == 1
    assert box_instances[0].windowTitle() == "检测到手写填空标记"

    assert window._doc_pages == [str(page_img)]
    assert window._regions == [region]
    assert "完整底图" not in window._ui.lineEdit_14.text()
    assert window._ui.lineEdit_2.text() == str(page_img)


def test_import_document_cancel(window, monkeypatch, tmp_path: Path) -> None:
    """底图导入：检测到标记时点击‘取消’，退出导入且不改变已有底图和区域。"""
    from handwritesim.core.models import TextRegion
    import handwritesim.core.doc_render as doc_render

    initial_bg = tmp_path / "initial_bg.png"
    initial_bg.touch()
    window._doc_pages = [str(initial_bg)]
    window._regions = []
    window._ui.lineEdit_2.setText(str(initial_bg))
    window._ui.lineEdit_14.setText("initial.pdf（1 页）")

    doc_file = tmp_path / "mock.pdf"
    doc_file.touch()
    page_img = tmp_path / "page_0.png"
    page_img.touch()
    region = TextRegion(
        x=10, y=10, w=100, h=40, page=1, role_id=2, highlight="yellow"
    )

    monkeypatch.setattr(
        QFileDialog,
        "getOpenFileName",
        lambda *args, **kwargs: (str(doc_file), "文档 (*.pdf *.docx)"),
    )
    monkeypatch.setattr(
        QMessageBox, "question", lambda *args, **kwargs: QMessageBox.StandardButton.Yes
    )
    monkeypatch.setattr(QMessageBox, "information", lambda *args, **kwargs: None)
    monkeypatch.setattr(QMessageBox, "warning", lambda *args, **kwargs: None)
    monkeypatch.setattr(
        doc_render,
        "document_to_page_images_with_regions",
        lambda path, out_dir, dpi=200: ([page_img], [region]),
    )

    box_instances = []

    def fake_exec(self):
        box_instances.append(self)
        for btn in self.buttons():
            if "取消" in btn.text():
                self._clicked_btn = btn
                break
        return 0

    monkeypatch.setattr(QMessageBox, "exec", fake_exec)
    monkeypatch.setattr(
        QMessageBox, "clickedButton", lambda self: getattr(self, "_clicked_btn", None)
    )

    window._import_document()

    assert len(box_instances) == 1
    assert box_instances[0].windowTitle() == "检测到手写填空标记"
    assert window._doc_pages == [str(initial_bg)]
    assert window._regions == []
    assert window._ui.lineEdit_2.text() == str(initial_bg)
    assert window._ui.lineEdit_14.text() == "initial.pdf（1 页）"


def test_import_document_without_regions_no_dialog(
    window, monkeypatch, tmp_path: Path
) -> None:
    """底图导入：无填空标记时静默导入普通底图，不弹手写填空标记选择对话框。"""
    import handwritesim.core.doc_render as doc_render

    doc_file = tmp_path / "plain.pdf"
    doc_file.touch()
    page_img = tmp_path / "plain_0.png"
    page_img.touch()

    monkeypatch.setattr(
        QFileDialog,
        "getOpenFileName",
        lambda *args, **kwargs: (str(doc_file), "文档 (*.pdf *.docx)"),
    )
    monkeypatch.setattr(
        QMessageBox, "question", lambda *args, **kwargs: QMessageBox.StandardButton.Yes
    )
    monkeypatch.setattr(
        doc_render,
        "document_to_page_images_with_regions",
        lambda path, out_dir, dpi=200: ([page_img], []),
    )

    box_instances = []

    def fake_exec(self):
        box_instances.append(self)
        return 0

    monkeypatch.setattr(QMessageBox, "exec", fake_exec)
    monkeypatch.setattr(QMessageBox, "clickedButton", lambda self: None)

    window._import_document()

    highlight_dialogs = [
        b for b in box_instances if b.windowTitle() == "检测到手写填空标记"
    ]
    assert len(highlight_dialogs) == 0
    assert window._doc_pages == [str(page_img)]
    assert window._regions == []



