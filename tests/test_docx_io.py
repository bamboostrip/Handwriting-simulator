"""docx 解析测试。"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from handwritesim.core.docx_io import load_paragraphs
from handwritesim.core.models import Paragraph


def _make_docx(path) -> None:
    doc = Document()
    hp = doc.add_paragraph("会议通知")
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body = doc.add_paragraph()
    body.paragraph_format.first_line_indent = Pt(24)  # 2 字符 @12pt
    body.add_run("现将有关事项通知如下。")
    doc.save(path)


def test_load_paragraphs(tmp_path):
    docx_path = tmp_path / "test.docx"
    _make_docx(docx_path)
    paras = load_paragraphs(docx_path)
    assert paras[0].align == "center"
    assert paras[0].first_line_indent == 0
    assert paras[1].align == "left"
    assert paras[1].first_line_indent > 0
    assert isinstance(paras[1], Paragraph)


def test_right_align(tmp_path):
    doc = Document()
    p = doc.add_paragraph("汇报人：张三")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    path = tmp_path / "right.docx"
    doc.save(path)
    paras = load_paragraphs(path)
    assert paras[0].align == "right"


def test_first_line_chars_indent(tmp_path):
    """中文 Word 的“首行缩进 2 字符”写 firstLineChars，应按字号换算。"""
    doc = Document()
    p = doc.add_paragraph("现将有关事项通知如下。")
    pPr = p._p.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn("w:ind"), {qn("w:firstLineChars"): "200"}))
    path = tmp_path / "chars.docx"
    doc.save(path)
    paras = load_paragraphs(path, font_size=140)
    assert paras[0].first_line_indent == 280


def test_style_inherited_indent(tmp_path):
    """直接格式缺失时应沿样式链继承首行缩进。"""
    doc = Document()
    doc.styles["Normal"].paragraph_format.first_line_indent = Pt(24)
    doc.add_paragraph("继承缩进的正文。")
    path = tmp_path / "style.docx"
    doc.save(path)
    paras = load_paragraphs(path)
    assert paras[0].first_line_indent > 0


def test_first_line_twips_follows_doc_font_size(tmp_path):
    """Word/WPS 把“首行缩进 2 字符”存成固定值 firstLine（480twips=12pt 的 2 字符）
    时不写 firstLineChars，应按文档字号还原为字符数再随渲染字号缩放。

    回归：曾按 96dpi 固定换算成 32px，渲染字号大时缩进明显偏小，
    用户二季度文档（仅 firstLine=480）导入后看起来没有首行缩进。
    """
    doc = Document()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)  # 等价 firstLine=480 twips
    p.add_run("现将有关事项通知如下。").font.size = Pt(12)  # 文档字号小四
    path = tmp_path / "twips.docx"
    doc.save(path)
    paras = load_paragraphs(path, font_size=140)
    assert paras[0].first_line_indent == 280  # 2 字符 × 140


def test_first_line_twips_fallback_style_font(tmp_path):
    """段落无 run 字号时，按样式/Normal 字号换算（24pt / 12pt = 2 字符）。"""
    doc = Document()
    doc.styles["Normal"].font.size = Pt(12)
    p = doc.add_paragraph("现将有关事项通知如下。")
    p.paragraph_format.first_line_indent = Pt(24)
    path = tmp_path / "twips_style.docx"
    doc.save(path)
    paras = load_paragraphs(path, font_size=140)
    assert paras[0].first_line_indent == 280


def test_first_line_chars_takes_priority(tmp_path):
    """存在 firstLineChars 时优先按字符数，忽略 firstLine 固定值。"""
    doc = Document()
    p = doc.add_paragraph("现将有关事项通知如下。")
    p.paragraph_format.first_line_indent = Pt(24)  # 固定值分支会算出不同结果
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = pPr.makeelement(qn("w:ind"), {})
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), "200")
    path = tmp_path / "both.docx"
    doc.save(path)
    paras = load_paragraphs(path, font_size=140)
    assert paras[0].first_line_indent == 280


def test_has_docx_highlights_and_ignore(tmp_path):
    from handwritesim.core.docx_io import has_docx_highlights, load_paragraphs_with_runs
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    doc_path = tmp_path / "test_hl.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Normal text ")
    r2 = p.add_run("Highlighted text")
    r2.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.save(str(doc_path))

    assert has_docx_highlights(doc_path) is True

    # 默认模式：混排
    paras_mixed = load_paragraphs_with_runs(doc_path, 36, ignore_highlights=False)
    roles_mixed = {r.role_id for p in paras_mixed for r in p.runs}
    assert 1 in roles_mixed  # 包含打印体
    assert any(rid >= 2 for rid in roles_mixed)  # 包含高亮角色

    # 忽略高亮模式：全部手写
    paras_hand = load_paragraphs_with_runs(doc_path, 36, ignore_highlights=True)
    roles_hand = {r.role_id for p in paras_hand for r in p.runs}
    assert roles_hand == {0}  # 只有默认手写

    # 无高亮文档与异常路径测试
    no_hl_path = tmp_path / "no_hl.docx"
    doc2 = Document()
    doc2.add_paragraph("Plain text only")
    doc2.save(str(no_hl_path))
    assert has_docx_highlights(no_hl_path) is False
    assert has_docx_highlights(tmp_path / "not_found.docx") is False
